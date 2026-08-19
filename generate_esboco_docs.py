import os
import json
import base64
import time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_docx():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header text
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Gerenciamento de Comandas — Documentação de Arquitetura da Informação"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(148, 163, 184)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Gerenciamento de Comandas")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(201, 42, 42) # #c92a2a

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Documentação de Arquitetura da Informação — Esboço v1")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_b = p.add_run(bold_prefix + " ")
            r_b.bold = True
            r_b.font.color.rgb = RGBColor(30, 41, 59)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    # 1. Objetivo
    add_section_header("1. Objetivo")
    p = doc.add_paragraph()
    p.add_run("Transformar a gestão tradicional de comandas, mesas, pedidos de cozinha e controle de fiados (antes feitos em blocos de papel e cadernos físicos) em uma aplicação web/mobile integrada, intuitiva e em tempo real, mantendo a agilidade necessária para o ambiente dinâmico de bares e restaurantes.")
    p = doc.add_paragraph()
    r = p.add_run("Problema central: ")
    r.bold = True
    p.add_run("extravio de comandas físicas, comandas ilegíveis, descompasso entre atendimento e cozinha (gargalos de produção), descontrole de limites de fiado e falta de visibilidade para o cliente acompanhar seus pedidos e saldo em tempo real.")

    # 2. Perfis de Uso
    add_section_header("2. Perfis de Uso")
    table_p = doc.add_table(rows=3, cols=2)
    table_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_p.autofit = False

    headers_p = ["Perfil", "Necessidade principal"]
    for i, h in enumerate(headers_p):
        cell = table_p.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell)

    rows_p = [
        ("Cliente (Autoatendimento)", "Acessar cardápio digital com fotos, montar pedido na cesta, enviar para produção, acompanhar o status ao vivo (Recebido ➔ Na Cozinha ➔ Pronto ➔ Entregue), consultar extrato e limite de fiado disponível."),
        ("Dono / Operador (Bistrô Gestão)", "Visualizar e gerenciar mesas/balcão em tempo real, lançar itens na comanda, controlar a fila de pedidos da cozinha (KDS), administrar o cardápio e preços, gerenciar clientes, definir limites de crédito e quitar fiados no caixa.")
    ]

    for row_idx, (perfil, nec) in enumerate(rows_p, start=1):
        c0 = table_p.cell(row_idx, 0)
        c0.text = perfil
        c0.paragraphs[0].runs[0].bold = True
        c1 = table_p.cell(row_idx, 1)
        c1.text = nec
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0)
        set_cell_margins(c1)

    doc.add_paragraph("Um único app, visão condicionada por login.").runs[0].font.italic = True

    # 3. Inventário de Conteúdo
    add_section_header("3. Inventário de Conteúdo")
    p = doc.add_paragraph("Itens levantados a partir do ambiente de restaurante e caderno de fiado:")
    add_bullet("Número da mesa, comanda avulsa (Balcão), quantidade de pessoas, status da mesa (Livre, Ocupada, Pagando, Reservada), horário de abertura.", "● Mesas e Comandas:")
    add_bullet("Itens selecionados, controle de quantidade (+/-), preço unitário, total da cesta, envio direto para produção.", "● Cesta de Pedidos:")
    add_bullet("Histórico de itens consumidos, subtotal, taxa de serviço (10% opcional), desconto em R$ e valor total.", "● Pedidos Confirmados:")
    add_bullet("Nome, descrição, categoria (Entradas, Pratos, Porções, Bebidas, Doces), preço e foto ilustrativa.", "● Cardápio / Catálogo:")
    add_bullet("Senha/Ticket, cliente, destino (Mesa ou Balcão), itens solicitados, observações, horário e status de preparo.", "● Fila da Cozinha (KDS):")
    add_bullet("Nome do cliente, telefone, saldo devedor acumulado, limite de crédito, status de risco e histórico.", "● Ficha de Clientes & Fiados:")
    add_bullet("Modal de pagamento com valor da dívida e forma de recebimento (Dinheiro, PIX, Cartão).", "● Quitação de Fiados:")
    add_bullet("Visualização do ticket atual, timeline de 4 etapas e aviso comemorativo de retirada de pedido.", "● Acompanhamento ao Vivo:")

    # 4. Card Sorting
    add_section_header("4. Card Sorting — Agrupamento")
    p = doc.add_paragraph("Ao organizar os itens acima a partir da jornada do cliente e do operador, os dados distribuem-se naturalmente em dois grandes grupos:")
    add_bullet("Cardápio ilustrado, cesta de compras, comprovante/senha gerada, timeline de acompanhamento ao vivo, carteira de fiado e limite de crédito.", "Grupo Cliente:")
    add_bullet("Mapa de mesas, gaveta de comandas, fila KDS da cozinha, editor de cardápio, ficha de clientes, limites de crédito e quitação de dívidas.", "Grupo Dono:")

    p_note = doc.add_paragraph()
    r = p_note.add_run("Nota: essa separação garante que o cliente interaja com uma interface leve e focada na experiência de consumo, enquanto a equipe do bistrô opera um painel de alta densidade informativa com controle total de fluxo.")
    r.font.italic = True
    r.font.color.rgb = RGBColor(100, 116, 139)

    # 5. Rotulagem
    add_section_header("5. Rotulagem")
    table_r = doc.add_table(rows=10, cols=2)
    table_r.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_r = ["Rótulo técnico / Sistema", "Rótulo usado no app"]
    for i, h in enumerate(headers_r):
        cell = table_r.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell)

    rows_r = [
        ("Perfil consumidor / Self-service", "Autoatendimento do Cliente"),
        ("Perfil administrador / Backoffice", "Bistrô & Gestão (Dono)"),
        ("Kitchen Display System (KDS)", "Gestão de Pedidos (Cozinha)"),
        ("Table Grid / Overview de Salão", "Visão das Mesas"),
        ("Drawer de Comanda / Order Slip", "Detalhes da Comanda"),
        ("Carrinho de Pedidos Pendentes", "Cesta de Pedidos"),
        ("Receivables / Contas a Receber", "Clientes & Fiados"),
        ("Payment Settlement / Baixa", "Quitar Dívida"),
        ("Order Live Tracker / Stepper", "Acompanhar Pedido")
    ]

    for row_idx, (tec, app_r) in enumerate(rows_r, start=1):
        c0 = table_r.cell(row_idx, 0)
        c0.text = tec
        c1 = table_r.cell(row_idx, 1)
        c1.text = app_r
        c1.paragraphs[0].runs[0].bold = True
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0)
        set_cell_margins(c1)

    p_crit = doc.add_paragraph("Critério: usar o vocabulário prático de restaurante, nunca termo técnico de sistema.")
    p_crit.runs[0].font.italic = True

    # 6. Organograma
    add_section_header("6. Organograma (Arquitetura da Informação)")
    tree_text = """LOGIN / ACESSO
 │
 ┌───────────────────────────┴───────────────────────────┐
 │                                                       │
PERFIL CLIENTE (Autoatendimento)                PERFIL DONO (Bistrô Gestão)
 │                                                       │
 ├─► Cardápio Ilustrado (Fotos & Categorias)             ├─► Visão das Mesas (Salão & Balcão)
 │     └─► Cesta de Pedidos (Ajuste de Quantidade)       │     ├─► Abertura de Comanda
 │                                                       │     ├─► Cesta de Lançamento Rápido
 ├─► Checkout (Mesa / Balcão / Viagem)                   │     ├─► Pedidos Confirmados (Subtotal/Taxa)
 │     └─► Pagamento (PIX, Cartão, Fiado com Limite)     │     └─► Fechamento & Pagamento
 │                                                       │
 ├─► Comprovante / Senha Gerada (Ticket)                 ├─► Gestão de Pedidos (KDS Cozinha)
 │                                                       │     ├─► Em Preparo (Filtro por Destino)
 └─► Acompanhar Pedido (Live Tracker em 4 passos)        │     ├─► Pronto (Notificação de Retirada)
       ├─► 1. Pedido Recebido                            │     └─► Concluído / Entregue
       ├─► 2. Na Cozinha (Em Preparo)                    │
       ├─► 3. Pronto para Retirada                       ├─► Cardápio / Preços (CRUD & Categorias)
       └─► 4. Entregue / Finalizado                      │
                                                         ├─► Clientes & Fiados
                                                         │     ├─► Consulta & Filtro em Tempo Real
                                                         │     ├─► Limite de Crédito & Saldo Devedor
                                                         │     └─► Modal de Quitação de Dívida
                                                         │
                                                         └─► Configurações & Métricas de Vendas"""
    p_tree = doc.add_paragraph()
    r_tree = p_tree.add_run(tree_text)
    r_tree.font.name = 'Courier New'
    r_tree.font.size = Pt(8.5)
    r_tree.font.color.rgb = RGBColor(15, 23, 42)

    # 7. Fluxogramas de Tarefa
    add_section_header("7. Fluxogramas de Tarefa")
    doc.add_paragraph("7.1 Cliente realiza pedido e acompanha em tempo real").runs[0].bold = True
    p = doc.add_paragraph("Login Cliente ➔ Cardápio ➔ Adiciona à Cesta ➔ Finaliza Pedido ➔ Escolhe Pagamento ➔ Recebe Senha ➔ Acompanha Timeline ao Vivo ➔ Retira o Pedido")
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    doc.add_paragraph("7.2 Dono abre mesa, lança itens e fecha conta no salão").runs[0].bold = True
    p = doc.add_paragraph("Visão das Mesas ➔ Clica em Mesa Livre ➔ Informa Cliente e Pessoas ➔ Abre Comanda ➔ Adiciona Produtos da Cesta ➔ Envia para Produção ➔ Aplica Taxa/Desconto ➔ Encerra Conta & Registra Pagamento")
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    doc.add_paragraph("7.3 Cliente compra no Fiado com validação de limite").runs[0].bold = True
    p = doc.add_paragraph("Checkout ➔ Seleciona 'Pagar Depois (Fiado)' ➔ Sistema verifica (Dívida + Pedido) vs Limite ➔ [Se Exceder: Bloqueia e exibe modal de alerta] ➔ [Se Permitido: Confirma pedido e lança débito na ficha]")
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    doc.add_paragraph("7.4 Dono quita dívida de cliente no caixa").runs[0].bold = True
    p = doc.add_paragraph("Menu Clientes & Fiados ➔ Busca Cliente ➔ Clica em 'Quitar Dívida' ➔ Abre Modal com Total Devedor ➔ Seleciona Forma de Recebimento (Dinheiro/PIX/Cartão) ➔ Confirma Baixa ➔ Saldo Zerado e Notificação de Sucesso")
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    # 8. Navegação e Sistema de Busca
    add_section_header("8. Navegação e Sistema de Busca")
    add_bullet("Barra lateral esquerda fixa (Sidebar) com ícones e rótulos claros para alternância rápida entre Visão das Mesas, Gestão de Pedidos (com contador em tempo real), Cardápio, Clientes & Fiados e Configurações.", "● Painel do Dono:")
    add_bullet("Barra de passos (Stepper) intuitiva com indicador de progresso (1. Cardápio ➔ 2. Cesta ➔ 3. Identificação ➔ 4. Pagamento ➔ 5. Pedido Confirmado / Acompanhamento).", "● Autoatendimento do Cliente:")
    add_bullet("Campo de busca em tempo real com filtro simultâneo por categorias (Todos, Entradas, Pratos, Porções, Bebidas, Doces) a cada tecla digitada.", "● Busca no Cardápio:")
    add_bullet("Busca instantânea por nome ou telefone com ordenação por maior saldo devedor e indicador visual de limite atingido.", "● Busca de Clientes:")

    # 9. Wireframes
    add_section_header("9. Wireframes — Especificação de Baixa Fidelidade")
    doc.add_paragraph("Estrutura de conteúdo por tela, sem estilo visual — pronto para prototipação em Figma/papel:").runs[0].font.italic = True
    
    add_bullet("Cards de mesas coloridos por estado (Livre, Ocupada, Pagando, Reservada), faturamento do dia e atalho de Venda Rápida.", "Home / Visão das Mesas (Dono):")
    add_bullet("Cesta de itens pendentes com botões (+/-), Pedidos confirmados com subtotal e taxa 10%, Seletor de cardápio e botão de fechar comanda.", "Drawer de Comanda (Dono):")
    add_bullet("Filtro por destino, cards com senha, cliente, tempo decorrido, itens com observações e botões de avanço de etapa.", "Gestão de Pedidos / KDS (Cozinha):")
    add_bullet("Resumo financeiro no topo (Total Fiado a Receber, Limite Concedido), tabela com saldo devedor e botão 'Quitar Dívida'.", "Clientes & Fiados (Dono):")
    add_bullet("Boas-vindas, campo de nome do cliente e seleção de mesa ou balcão.", "Identificação / Login (Cliente):")
    add_bullet("Abas de categoria, cards de produtos com foto, nome, descrição, preço e botão 'Adicionar à Cesta'.", "Cardápio do Cliente:")
    add_bullet("Senha em destaque, estimativa de tempo e timeline visual com os 4 passos de preparo.", "Acompanhar Pedido / Live Tracker (Cliente):")

    # 10. Prototipagem e Escala de Fidelidade
    add_section_header("10. Prototipagem e Escala de Fidelidade")
    table_f = doc.add_table(rows=4, cols=4)
    table_f.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_f = ["Estágio", "O que existe", "O que se testa", "Ferramenta"]
    for i, h in enumerate(headers_f):
        cell = table_f.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell)

    rows_f = [
        ("Lo-fi", "Blocos, caixas cinzas, rótulos, hierarquia básica", "Se a arquitetura faz sentido", "Papel ou Figma em cinza"),
        ("Mid-fi", "Layout navegável real, componentes sem estilo final", "Se o fluxo e comanda funcionam", "Figma protótipo / HTML"),
        ("Hi-fi", "Cores finais, tipografia, fotos reais, microtextos e sincronização", "Se a interface comunica confiança e apetite", "Web App funcional (HTML/CSS/JS)")
    ]

    for row_idx, (est, ex, test, ferr) in enumerate(rows_f, start=1):
        c0 = table_f.cell(row_idx, 0)
        c0.text = est
        c0.paragraphs[0].runs[0].bold = True
        c1 = table_f.cell(row_idx, 1)
        c1.text = ex
        c2 = table_f.cell(row_idx, 2)
        c2.text = test
        c3 = table_f.cell(row_idx, 3)
        c3.text = ferr
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for c in [c0, c1, c2, c3]:
            set_cell_background(c, bg)
            set_cell_margins(c)

    # 11. Design de Interação
    add_section_header("11. Design de Interação (IxD)")
    add_bullet("Abre em camada sobreposta (modal lateral) com fundo desfocado, preservando a visão geral do salão.", "Gaveta Lateral (Drawer de Comanda):")
    add_bullet("Cabeçalhos em acordeão com transição suave para expansão e recolhimento instantâneo.", "Cesta & Pedidos Confirmados:")
    add_bullet("Microfeedback visual ao toque para adicionar e diminuir itens sem digitação.", "Controles de Quantidade (+ / -):")
    add_bullet("Bloqueio visual do botão de confirmação e modal explicativo detalhado ao exceder o limite de crédito.", "Validação de Fiado:")
    add_bullet("Banners flutuantes com animação de subida para confirmações e baixas de pagamentos.", "Notificações Toast:")
    add_bullet("Sincronização em tempo real via storage events reflete mudanças de preparo na tela do cliente sem reload.", "Live Sync:")

    # 12. UI — Interface Final
    add_section_header("12. UI — Interface Final (Guia de Estilo)")
    table_ui = doc.add_table(rows=8, cols=3)
    table_ui.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_ui = ["Uso / Token", "Cor & HEX", "Justificativa & Aplicação"]
    for i, h in enumerate(headers_ui):
        cell = table_ui.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell)

    rows_ui = [
        ("Primária", "#C92A2A · Vermelho Bistrô", "Transmite apetite, acolhimento e autoridade gastronômica (botões principais, logotipo, destaques)."),
        ("Primária Dark / Hover", "#A61E1E · Rubi Intenso", "Estados de hover e gradientes sofisticados para botões e cabeçalhos."),
        ("Fundo Suave / Badge", "#FEE2E2 · Vermelho Claro", "Fundo de badges, alertas suaves e seleção de itens no cardápio."),
        ("Status Livre / Sucesso", "#10B981 · Verde Esmeralda", "Mesas livres, confirmação de pedidos e pedidos prontos para entrega."),
        ("Status Ocupada", "#F59E0B · Âmbar / Laranja", "Mesas ocupadas e itens em preparo na cozinha."),
        ("Fundo da Aplicação", "#EAEDF2 · Branco Acinzentado", "Contraste perfeito com as caixas brancas dos cards, evitando cansaço visual."),
        ("Texto / Tipografia", "#1E293B · Grafite Escuro", "Altíssima legibilidade em telas de balcão e smartphones sob qualquer iluminação.")
    ]

    for row_idx, (uso, cor, just) in enumerate(rows_ui, start=1):
        c0 = table_ui.cell(row_idx, 0)
        c0.text = uso
        c0.paragraphs[0].runs[0].bold = True
        c1 = table_ui.cell(row_idx, 1)
        c1.text = cor
        c2 = table_ui.cell(row_idx, 2)
        c2.text = just
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for c in [c0, c1, c2, c3 if 'c3' in locals() else c2]:
            set_cell_background(c, bg)
            set_cell_margins(c)

    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    add_bullet("Outfit (títulos e marca) + Plus Jakarta Sans (textos e números tabulares de comanda).", "● Tipografia Oficial:")
    add_bullet("Cards de mesas com borda de status, Drawer lateral retrátil, KDS cards com contador de tempo e Timeline de 4 etapas.", "● Componentes-chave:")

    docx_path = r"c:\Users\aluno\Downloads\Emerson\Gerenciamento-de-Comandas-main\Gerenciamento-de-Comandas-main\documentacao-esboco-gerenciamento-de-comandas.docx"
    doc.save(docx_path)
    print(f"DOCX saved successfully at {docx_path}")

def generate_pdf():
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    
    html_path = os.path.abspath(r"c:\Users\aluno\Downloads\Emerson\Gerenciamento-de-Comandas-main\Gerenciamento-de-Comandas-main\documentacao-esboco-gerenciamento-de-comandas.html")
    pdf_path = os.path.abspath(r"c:\Users\aluno\Downloads\Emerson\Gerenciamento-de-Comandas-main\Gerenciamento-de-Comandas-main\documentacao-esboco-gerenciamento-de-comandas.pdf")
    
    options = Options()
    options.add_argument('--headless')
    options.add_argument('--disable-gpu')
    options.add_argument('--no-sandbox')
    
    driver = webdriver.Chrome(options=options)
    driver.get(f"file:///{html_path}")
    
    # Wait for Mermaid to render
    time.sleep(3)
    
    print_options = {
        'landscape': False,
        'displayHeaderFooter': False,
        'printBackground': True,
        'preferCSSPageSize': True,
        'marginTop': 0,
        'marginBottom': 0,
        'marginLeft': 0,
        'marginRight': 0,
    }
    
    result = driver.execute_cdp_cmd("Page.printToPDF", print_options)
    with open(pdf_path, 'wb') as f:
        f.write(base64.b64decode(result['data']))
        
    driver.quit()
    print(f"PDF saved successfully at {pdf_path}")

if __name__ == '__main__':
    generate_docx()
    generate_pdf()
