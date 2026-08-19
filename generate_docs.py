# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Cores do Sistema
    COLOR_PRIMARY = RGBColor(99, 102, 241)     # #6366f1 (Indigo)
    COLOR_SECONDARY = RGBColor(16, 185, 129)   # #10b981 (Emerald)
    COLOR_DARK = RGBColor(15, 23, 42)          # #0f172a (Slate)
    COLOR_MUTED = RGBColor(100, 116, 139)      # #64748b

    # Título Principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("BISTRÔ GESTÃO & AUTOATENDIMENTO")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("Manual de Operação e Fluxogramas do Sistema (Dono & Cliente)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SEÇÃO 1: VISÃO GERAL ---
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Visão Geral da Arquitetura Integrada")
    r1.font.name = 'Arial'
    r1.font.color.rgb = COLOR_PRIMARY

    p_desc = doc.add_paragraph(
        "O sistema é dividido em dois módulos totalmente sincronizados em tempo real via navegador e armazenamento local:\n"
        "• Módulo do Dono/Gestor (index.html): Painel administrativo para gestão de mesas, pedidos em tempo real (KDS), cardápio, controle de fiados e receita.\n"
        "• Módulo do Cliente (cliente.html): Portal responsivo para autoatendimento em totens, tablets ou celulares com consulta de saldo, login de clientes cadastrados, cardápio digital e emissão de senhas de balcão."
    )
    p_desc.style.font.name = 'Arial'
    p_desc.style.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # --- SEÇÃO 2: FLUXO DO DONO / GESTOR ---
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Fluxograma do Dono / Gestor (index.html)")
    r2.font.name = 'Arial'
    r2.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "O painel de gestão centraliza todas as operações de salão, balcão e retaguarda. Abaixo está o fluxo passo a passo de cada etapa:"
    )

    # Tabela do Dono
    table_dono = doc.add_table(rows=1, cols=3)
    table_dono.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_dono.autofit = False

    hdr_cells = table_dono.rows[0].cells
    headers = ["Módulo / Tela", "Ações do Gestor", "Resultado no Sistema"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "4F46E5")
        set_cell_margins(hdr_cells[i], 120, 120, 140, 140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

    passos_dono = [
        ("1. Visão das Mesas", "Clica em uma mesa livre para abri-la informando o nome do cliente e quantidade de pessoas.", "A mesa muda de status para 'Ocupada' e inicia a contagem de tempo de permanência."),
        ("2. Gaveta de Comandas", "Adiciona pratos e bebidas à cesta da mesa e clica em 'Confirmar e Enviar para Produção'.", "Os itens entram na comanda da mesa e são disparados instantaneamente para a fila de produção da cozinha."),
        ("3. Fechamento de Conta", "Abre a gaveta da mesa, aplica taxa de serviço (10%) ou desconto e seleciona o método de pagamento.", "Calcula subtotal e total com recibo detalhado."),
        ("4. Pagamento & Liberação", "Recebe via Pix, Cartão, Dinheiro ou Fiado e clica em 'Confirmar Pagamento e Liberar'.", "Soma à receita do dia (ou lança na lista de fiados), limpa o consumo e libera a mesa para novos clientes."),
        ("5. Gestão de Pedidos (KDS)", "Acompanha os pedidos das mesas e do autoatendimento em cartões coloridos organizados por horário.", "Pode alterar status para '🟡 Em Preparo', '🟢 Pronto para Entrega' ou '⚪ Entregue'."),
        ("6. Cardápio & Preços", "Cadastra novos pratos com foto, categoria, descrição e preço, ou edita itens existentes.", "Atualiza instantaneamente o cardápio no painel e na tela de autoatendimento dos clientes."),
        ("7. Clientes & Fiados", "Cadastra clientes com limite de crédito e dia de vencimento; consulta extrato detalhado de itens consumidos.", "Permite imprimir nota promissória para assinatura ou quitar a dívida após recebimento.")
    ]

    for row_idx, (mod, acao, res) in enumerate(passos_dono):
        row = table_dono.add_row()
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for c_idx, text in enumerate([mod, acao, res]):
            cell = row.cells[c_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 100, 100, 120, 120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if c_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Diagrama textual do Dono
    doc.add_heading("Diagrama Sequencial do Dono:", level=2).runs[0].font.size = Pt(12)
    diag_dono = doc.add_paragraph(
        "[ Início: Login no Painel ] ➔ [ Visão das Mesas / Balcão ]\n"
        "   ├─➔ [ Abrir Mesa / Comanda ] ➔ [ Lançar Pratos ] ➔ [ Enviar Cozinha ]\n"
        "   ├─➔ [ Gestão de Pedidos ] ➔ [ Preparo 🟡 ] ➔ [ Pronto 🟢 ] ➔ [ Entregue ⚪ ]\n"
        "   ├─➔ [ Fechamento ] ➔ [ Pix / Cartão / Dinheiro ] ➔ [ Receita do Dia ] ➔ [ Liberar Mesa ]\n"
        "   ├─➔ [ Opção Fiado ] ➔ [ Lançar na Caderneta ] ➔ [ Imprimir Nota ] ➔ [ Quitar Débito ]\n"
        "   └─➔ [ Cardápio / Preços ] ➔ [ Atualizar Preços e Fotos ] ➔ [ Sincronização Geral ]"
    )
    diag_dono.style.font.name = 'Consolas'
    diag_dono.style.font.size = Pt(9)
    diag_dono.paragraph_format.left_indent = Inches(0.2)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # --- SEÇÃO 3: FLUXO DO CLIENTE ---
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. Fluxograma do Cliente (cliente.html)")
    r3.font.name = 'Arial'
    r3.font.color.rgb = COLOR_SECONDARY

    doc.add_paragraph(
        "A área do cliente foi projetada com alta usabilidade e responsividade para autoatendimento ágil no balcão e totens:"
    )

    # Tabela do Cliente
    table_cli = doc.add_table(rows=1, cols=3)
    table_cli.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cli.autofit = False

    hdr_cells_cli = table_cli.rows[0].cells
    for i, h in enumerate(["Etapa do Cliente", "Ação Realizada", "Comportamento do Sistema"]):
        hdr_cells_cli[i].text = h
        set_cell_background(hdr_cells_cli[i], "059669")
        set_cell_margins(hdr_cells_cli[i], 120, 120, 140, 140)
        p = hdr_cells_cli[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

    passos_cli = [
        ("1. Tela Inicial", "O cliente informa seu Nome, Telefone e opção de consumo (Para Viagem ou Comer no Local).", "Pode avançar diretamente para o cardápio ou clicar em 'Fazer Login / Minha Conta'."),
        ("2. Login (Cadastrado)", "Seleciona seu nome na lista de clientes cadastrados e digita sua senha de acesso.", "Autentica o cliente e abre seu painel personalizado com saldo e extrato."),
        ("3. Minha Conta / Fiado", "Consulta o Total Devido Atual, Limite Disponível, Data de Vencimento e itens que consumiu no passado.", "Oferece botões para 'Fazer Pedido no Fiado' ou 'Pagar Conta via PIX' (com QR Code dinâmico)."),
        ("4. Cardápio Interativo", "Navega pelas categorias (Entradas, Pratos, Bebidas, Doces) ou pesquisa pratos por nome.", "Exibe fotos, descrições, preços e botões '+' e '−' com barra flutuante inferior de resumo em tempo real."),
        ("5. Carrinho & Revisão", "Ajusta quantidades dos pratos e adiciona observações personalizadas (ex: 'Sem cebola').", "Calcula subtotal e prepara para a tela de pagamento."),
        ("6. Pagamento no Balcão", "Escolhe entre PIX (QR Code e Copia e Cola), Cartão no Balcão, Dinheiro (com cálculo de troco) ou Fiado.", "Se escolher Fiado, vincula à caderneta e aprova a compra."),
        ("7. Senha & Confirmação", "O sistema emite a Senha de Retirada (ex: Senha #001) com resumo dos itens e horário.", "Dispara o pedido imediatamente para o painel de produção da cozinha e permite imprimir cupom térmico.")
    ]

    for row_idx, (etapa, acao, comp) in enumerate(passos_cli):
        row = table_cli.add_row()
        bg_color = "F0FDF4" if row_idx % 2 == 0 else "FFFFFF"
        for c_idx, text in enumerate([etapa, acao, comp]):
            cell = row.cells[c_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 100, 100, 120, 120)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if c_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = COLOR_SECONDARY

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Diagrama textual do Cliente
    doc.add_heading("Diagrama Sequencial do Cliente:", level=2).runs[0].font.size = Pt(12)
    diag_cli = doc.add_paragraph(
        "[ Início: cliente.html ]\n"
        "   ├─➔ [ Opção 1: Pedido Rápido ] ➔ [ Nome + Opção de Consumo ] ➔ [ Cardápio ]\n"
        "   └─➔ [ Opção 2: Login Cadastrado ] ➔ [ Nome + Senha ]\n"
        "          ├─➔ [ Minha Conta / Fiado ] ➔ [ Consultar Extrato / Limite ] ➔ [ Pagar PIX ]\n"
        "          └─➔ [ Fazer Pedido ] ➔ [ Cardápio Interativo com Fotos ]\n"
        "                 ➔ [ Carrinho + Observações para Cozinha ]\n"
        "                 ➔ [ Pagamento: PIX / Cartão / Dinheiro / Fiado ]\n"
        "                 ➔ [ Confirmação: Emissão de Senha #00X + Envio Cozinha ]"
    )
    diag_cli.style.font.name = 'Consolas'
    diag_cli.style.font.size = Pt(9)
    diag_cli.paragraph_format.left_indent = Inches(0.2)

    # Salvar o documento Word
    docx_path = "Fluxograma_e_Manual_do_Sistema.docx"
    doc.save(docx_path)
    print(f"Documento Word criado com sucesso em: {docx_path}")

if __name__ == '__main__':
    create_document()
